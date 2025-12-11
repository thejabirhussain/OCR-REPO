"""PDF and DOCX extraction service with structure preservation."""

import logging
from pathlib import Path
from typing import List, Optional

import pdfplumber
import pymupdf
from docx import Document as DocxDocument

from app.utils.layout_schema import create_block
from app.schemas import Page

logger = logging.getLogger(__name__)


def is_text_based_pdf(pdf_path: Path) -> bool:
    """Check if PDF has extractable text layer."""
    try:
        doc = pymupdf.open(pdf_path)
        has_text = False
        for page_num in range(min(3, len(doc))):  # Check first 3 pages
            page = doc[page_num]
            text = page.get_text().strip()
            if len(text) > 50:  # Reasonable amount of text
                has_text = True
                break
        doc.close()
        return has_text
    except Exception as e:
        logger.error(f"Error checking PDF text layer: {e}")
        return False


def extract_text_from_pdf_pymupdf(
    pdf_path: Path,
    ocr_callback=None,
    ocr_engine: str = "paddleocr",
) -> List[dict]:
    """
    Extract text from PDF using PyMuPDF with structure preservation.

    Args:
        pdf_path: Path to PDF file
        ocr_callback: Function to call for OCR on image pages: (image_path, page_index) -> blocks
        ocr_engine: OCR engine to use if OCR is needed

    Returns:
        List of page dictionaries
    """
    pages = []
    doc = pymupdf.open(pdf_path)

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = []

            # Try to extract text directly
            text_dict = page.get_text("dict")
            text_content = page.get_text().strip()

            if len(text_content) > 10:
                # Text-based page
                block_id_counter = 0
                for block in text_dict["blocks"]:
                    if "lines" in block:  # Text block
                        block_text_parts = []
                        for line in block["lines"]:
                            line_text_parts = []
                            for span in line["spans"]:
                                line_text_parts.append(span["text"])
                            if line_text_parts:
                                block_text_parts.append(" ".join(line_text_parts))

                        if block_text_parts:
                            block_text = "\n".join(block_text_parts)
                            bbox = block["bbox"]  # [x1, y1, x2, y2]

                            # Detect if heading (simple heuristic: larger font or at top)
                            is_heading = False
                            heading_level = None
                            if "lines" in block and block["lines"]:
                                first_span = block["lines"][0]["spans"][0]
                                font_size = first_span.get("size", 12)
                                if font_size > 14 or bbox[1] < 100:  # Top of page or large font
                                    is_heading = True
                                    heading_level = 1 if font_size > 16 else 2

                            block_obj = create_block(
                                block_id=f"{page_num}-{block_id_counter}",
                                text=block_text,
                                block_type="heading" if is_heading else "paragraph",
                                bbox=list(bbox),
                                is_heading=is_heading,
                                heading_level=heading_level,
                            )
                            blocks.append(block_obj.dict())
                            block_id_counter += 1
            else:
                # Image-based page - use OCR
                if ocr_callback:
                    logger.info(f"Page {page_num} appears to be image-based, using OCR")
                    # Convert page to image
                    pix = page.get_pixmap(matrix=pymupdf.Matrix(2, 2))  # 2x zoom for better OCR
                    img_path = pdf_path.parent / f"temp_page_{page_num}.png"
                    pix.save(str(img_path))
                    try:
                        ocr_blocks = ocr_callback(img_path, page_num, ocr_engine)
                        blocks.extend(ocr_blocks)
                    finally:
                        if img_path.exists():
                            img_path.unlink()
                else:
                    logger.warning(f"Page {page_num} has no text and OCR callback not provided")

            pages.append(Page(page_index=page_num, blocks=[create_block(**b) for b in blocks]).dict())

    finally:
        doc.close()

    return pages


def extract_text_from_pdf_pdfplumber(pdf_path: Path) -> List[dict]:
    """
    Extract text from PDF using pdfplumber (alternative method).

    Returns:
        List of page dictionaries
    """
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            blocks = []
            block_id_counter = 0

            # Extract text with layout
            words = page.extract_words()
            if words:
                # Group words into paragraphs (simple approach)
                current_paragraph = []
                current_y = None
                y_threshold = 10

                for word in words:
                    word_y = word["top"]
                    if current_y is None or abs(word_y - current_y) < y_threshold:
                        current_paragraph.append(word["text"])
                        current_y = word_y
                    else:
                        # New paragraph
                        if current_paragraph:
                            block_text = " ".join(current_paragraph)
                            block_obj = create_block(
                                block_id=f"{page_num}-{block_id_counter}",
                                text=block_text,
                                block_type="paragraph",
                            )
                            blocks.append(block_obj.dict())
                            block_id_counter += 1
                        current_paragraph = [word["text"]]
                        current_y = word_y

                # Don't forget last paragraph
                if current_paragraph:
                    block_text = " ".join(current_paragraph)
                    block_obj = create_block(
                        block_id=f"{page_num}-{block_id_counter}",
                        text=block_text,
                        block_type="paragraph",
                    )
                    blocks.append(block_obj.dict())

            pages.append(Page(page_index=page_num, blocks=[create_block(**b) for b in blocks]).dict())

    return pages


def extract_text_from_docx(docx_path: Path) -> List[dict]:
    """
    Extract text from DOCX file with structure preservation.

    Returns:
        List of page dictionaries (DOCX doesn't have pages, so we use a single "page")
    """
    doc = DocxDocument(docx_path)
    blocks = []
    block_id_counter = 0

    for element in doc.element.body:
        if element.tag.endswith("p"):  # Paragraph
            para = None
            for para_obj in doc.paragraphs:
                if para_obj._element == element:
                    para = para_obj
                    break

            if para:
                text = para.text.strip()
                if text:
                    # Check if heading
                    is_heading = para.style.name.startswith("Heading")
                    heading_level = None
                    if is_heading:
                        try:
                            heading_level = int(para.style.name.split()[-1])
                        except (ValueError, IndexError):
                            heading_level = 1

                    block_type = "heading" if is_heading else "paragraph"
                    block_obj = create_block(
                        block_id=f"0-{block_id_counter}",
                        text=text,
                        block_type=block_type,
                        is_heading=is_heading,
                        heading_level=heading_level,
                    )
                    blocks.append(block_obj.dict())
                    block_id_counter += 1

        elif element.tag.endswith("tbl"):  # Table
            # Find corresponding table object
            table = None
            for table_obj in doc.tables:
                if table_obj._element == element:
                    table = table_obj
                    break

            if table:
                table_id = f"table-{block_id_counter}"
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            block_obj = create_block(
                                block_id=f"0-{block_id_counter}",
                                text=cell_text,
                                block_type="table_cell",
                                table_row=row_idx,
                                table_col=col_idx,
                                table_id=table_id,
                            )
                            blocks.append(block_obj.dict())
                            block_id_counter += 1

    # DOCX doesn't have pages, so we return a single page
    return [Page(page_index=0, blocks=[create_block(**b) for b in blocks]).dict()]


def extract_text_from_file(
    file_path: Path,
    file_type: str,
    ocr_callback=None,
    ocr_engine: str = "paddleocr",
) -> List[dict]:
    """
    Extract text from file based on type.

    Args:
        file_path: Path to file
        file_type: File type ('pdf', 'docx', 'image')
        ocr_callback: Function for OCR: (image_path, page_index, engine) -> blocks
        ocr_engine: OCR engine to use

    Returns:
        List of page dictionaries
    """
    logger.info(f"Extracting text from {file_type} file: {file_path}")

    if file_type == "pdf":
        # Try PyMuPDF first (better structure preservation)
        try:
            return extract_text_from_pdf_pymupdf(file_path, ocr_callback, ocr_engine)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}, trying pdfplumber")
            return extract_text_from_pdf_pdfplumber(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type in ["image", "jpg", "jpeg", "png", "tiff", "tif"]:
        if ocr_callback:
            return [ocr_callback(file_path, 0, ocr_engine)]
        else:
            raise ValueError("OCR callback required for image files")
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

