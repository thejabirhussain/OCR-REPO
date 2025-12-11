"""Generate DOCX files from structured documents."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.schemas import Block, StructuredDocument

logger = logging.getLogger(__name__)


def generate_docx_from_document(
    document: StructuredDocument,
    output_path: Path,
    preserve_structure: bool = True,
) -> Path:
    """
    Generate a DOCX file from a structured document.

    Args:
        document: Structured document (Arabic or English)
        output_path: Path where DOCX will be saved
        preserve_structure: Whether to preserve headings, lists, tables

    Returns:
        Path to generated DOCX file
    """
    logger.info(f"Generating DOCX from document: {output_path}")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    for page in document.pages:
        for block in page.blocks:
            if not block.text.strip():
                continue

            if preserve_structure:
                # Handle different block types
                if block.type == "heading" and block.metadata.is_heading:
                    # Add heading
                    level = block.metadata.heading_level or 1
                    heading = doc.add_heading(block.text, level=level)
                    # Set RTL for Arabic if needed
                    if document.language == "ar":
                        paragraph_format = heading.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                elif block.type == "table_cell" and block.metadata.table:
                    # Handle tables - this is simplified
                    # In a full implementation, you'd need to reconstruct table structure
                    # For now, we'll add table cells as paragraphs with indentation
                    para = doc.add_paragraph(block.text)
                    if block.metadata.table.col:
                        para.paragraph_format.left_indent = Inches(block.metadata.table.col * 0.5)

                elif block.metadata.list_level is not None:
                    # Add list item
                    para = doc.add_paragraph(block.text, style="List Bullet")
                    para.paragraph_format.left_indent = Inches(block.metadata.list_level * 0.5)

                else:
                    # Regular paragraph
                    para = doc.add_paragraph(block.text)
                    # Set RTL for Arabic
                    if document.language == "ar":
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # Simple mode: just add paragraphs
                doc.add_paragraph(block.text)

        # Add page break between pages (except last)
        if page.page_index < len(document.pages) - 1:
            doc.add_page_break()

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    logger.info(f"DOCX generated successfully: {output_path}")
    return output_path


def generate_txt_from_document(
    document: StructuredDocument,
    output_path: Path,
    preserve_structure: bool = True,
) -> Path:
    """
    Generate a plain text file from a structured document.

    Args:
        document: Structured document
        output_path: Path where TXT will be saved
        preserve_structure: Whether to add formatting markers

    Returns:
        Path to generated TXT file
    """
    logger.info(f"Generating TXT from document: {output_path}")

    lines = []

    for page in document.pages:
        for block in page.blocks:
            if not block.text.strip():
                continue

            if preserve_structure:
                if block.type == "heading" and block.metadata.is_heading:
                    level = block.metadata.heading_level or 1
                    prefix = "#" * level + " "
                    lines.append(f"{prefix}{block.text}")
                elif block.metadata.list_level is not None:
                    prefix = "  " * block.metadata.list_level + "- "
                    lines.append(f"{prefix}{block.text}")
                else:
                    lines.append(block.text)
            else:
                lines.append(block.text)

        lines.append("")  # Blank line between pages

    # Write to file
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info(f"TXT generated successfully: {output_path}")
    return output_path

